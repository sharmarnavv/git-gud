"""
Document loaders for resume parsing.

This module provides document loading utilities for different file formats
including PDF, DOCX, and TXT files.
"""

import os
import mimetypes
from pathlib import Path
from typing import Dict, Any

from .resume_interfaces import DocumentLoaderInterface
from .resume_exceptions import DocumentLoadError, FileFormatError, TextExtractionError
from job_parser.logging_config import get_logger


class BaseDocumentLoader(DocumentLoaderInterface):
    """Base class for document loaders with common functionality."""
    
    def __init__(self):
        self.logger = get_logger(__name__)
        self.supported_extensions = set()
        self.supported_mimetypes = set()
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if file can be processed.
        
        Args:
            file_path: Path to the file to validate
            
        Returns:
            True if file can be processed, False otherwise
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return False
            
            # Check file size (max 10MB)
            file_size = os.path.getsize(file_path)
            if file_size > 10 * 1024 * 1024:  # 10MB limit
                return False
            
            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.supported_extensions:
                return False
            
            # Check MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type and mime_type not in self.supported_mimetypes:
                return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"File validation failed for {file_path}: {e}")
            return False


class PDFLoader(BaseDocumentLoader):
    """PDF document loader using PyPDF2 and pdfplumber."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.pdf'}
        self.supported_mimetypes = {'application/pdf'}
        
        # Try to import PDF libraries
        self._pypdf2_available = False
        self._pdfplumber_available = False
        
        try:
            import PyPDF2
            self._pypdf2_available = True
        except ImportError:
            self.logger.warning("PyPDF2 not available, PDF extraction may be limited")
        
        try:
            import pdfplumber
            self._pdfplumber_available = True
        except ImportError:
            self.logger.warning("pdfplumber not available, using fallback PDF extraction")
    
    def load_document(self, file_path: str) -> str:
        """Load and extract text from PDF document.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentLoadError: If document loading fails
        """
        if not self.validate_file(file_path):
            raise FileFormatError(f"Invalid PDF file: {file_path}")
        
        try:
            # Try pdfplumber first (better text extraction)
            if self._pdfplumber_available:
                return self._extract_with_pdfplumber(file_path)
            
            # Fallback to PyPDF2
            elif self._pypdf2_available:
                return self._extract_with_pypdf2(file_path)
            
            else:
                raise DocumentLoadError("No PDF extraction libraries available")
                
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            raise TextExtractionError(f"PDF text extraction failed: {e}", cause=e)
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber."""
        import pdfplumber
        
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return '\n'.join(text_content)
    
    def _extract_with_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2."""
        import PyPDF2
        
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return '\n'.join(text_content)


class DOCXLoader(BaseDocumentLoader):
    """DOCX document loader using python-docx."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.docx'}
        self.supported_mimetypes = {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        # Try to import docx library
        self._docx_available = False
        
        try:
            import docx
            self._docx_available = True
        except ImportError:
            self.logger.warning("python-docx not available, DOCX files cannot be processed")
    
    def load_document(self, file_path: str) -> str:
        """Load and extract text from DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentLoadError: If document loading fails
        """
        if not self.validate_file(file_path):
            raise FileFormatError(f"Invalid DOCX file: {file_path}")
        
        if not self._docx_available:
            raise DocumentLoadError("python-docx library not available")
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise TextExtractionError(f"DOCX text extraction failed: {e}", cause=e)


class TXTLoader(BaseDocumentLoader):
    """Plain text document loader."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.txt', '.text'}
        self.supported_mimetypes = {'text/plain'}
    
    def load_document(self, file_path: str) -> str:
        """Load and extract text from TXT document.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentLoadError: If document loading fails
        """
        if not self.validate_file(file_path):
            raise FileFormatError(f"Invalid TXT file: {file_path}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()
                
        except Exception as e:
            self.logger.error(f"Failed to read text file {file_path}: {e}")
            raise TextExtractionError(f"Text file reading failed: {e}", cause=e)


class DocumentLoaderFactory:
    """Factory class for creating appropriate document loaders."""
    
    def __init__(self):
        self.logger = get_logger(__name__)
        self._loaders = {
            '.pdf': PDFLoader,
            '.docx': DOCXLoader,
            '.txt': TXTLoader,
            '.text': TXTLoader
        }
    
    def get_loader(self, file_path: str) -> DocumentLoaderInterface:
        """Get appropriate loader for file type.
        
        Args:
            file_path: Path to the file
            
        Returns:
            DocumentLoaderInterface instance for the file type
            
        Raises:
            FileFormatError: If file format is not supported
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self._loaders:
            raise FileFormatError(f"Unsupported file format: {file_ext}")
        
        loader_class = self._loaders[file_ext]
        return loader_class()
    
    def detect_file_format(self, file_path: str) -> Dict[str, Any]:
        """Detect file format and return metadata.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file format information
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            mime_type, _ = mimetypes.guess_type(file_path)
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            
            return {
                'extension': file_ext,
                'mime_type': mime_type,
                'size_bytes': file_size,
                'size_mb': file_size / (1024 * 1024),
                'supported': file_ext in self._loaders,
                'loader_available': self._check_loader_availability(file_ext)
            }
            
        except Exception as e:
            self.logger.error(f"Failed to detect file format for {file_path}: {e}")
            return {
                'extension': 'unknown',
                'mime_type': None,
                'size_bytes': 0,
                'size_mb': 0,
                'supported': False,
                'loader_available': False,
                'error': str(e)
            }
    
    def _check_loader_availability(self, file_ext: str) -> bool:
        """Check if loader dependencies are available for file extension."""
        try:
            if file_ext == '.pdf':
                try:
                    import PyPDF2
                    return True
                except ImportError:
                    try:
                        import pdfplumber
                        return True
                    except ImportError:
                        return False
            
            elif file_ext == '.docx':
                try:
                    import docx
                    return True
                except ImportError:
                    return False
            
            elif file_ext in ['.txt', '.text']:
                return True  # No external dependencies needed
            
            return False
            
        except Exception:
            return False
    
    def get_supported_formats(self) -> Dict[str, Dict[str, Any]]:
        """Get information about all supported formats.
        
        Returns:
            Dictionary with format information
        """
        formats = {}
        
        for ext, loader_class in self._loaders.items():
            formats[ext] = {
                'loader_class': loader_class.__name__,
                'available': self._check_loader_availability(ext),
                'description': self._get_format_description(ext)
            }
        
        return formats
    
    def _get_format_description(self, file_ext: str) -> str:
        """Get description for file format."""
        descriptions = {
            '.pdf': 'Portable Document Format - requires PyPDF2 or pdfplumber',
            '.docx': 'Microsoft Word Document - requires python-docx',
            '.txt': 'Plain Text File - no additional dependencies',
            '.text': 'Plain Text File - no additional dependencies'
        }
        
        return descriptions.get(file_ext, 'Unknown format')